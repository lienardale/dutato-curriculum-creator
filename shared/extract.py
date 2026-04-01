"""
Extract text from PDF/DOCX files with structure preservation.

Outputs a JSON file with:
- pages: [{page: int, text: str}]
- toc: [{title, level, page}] (from PDF bookmarks)
- metadata: {title, author, total_pages}
"""

import json
import sys
from pathlib import Path

import fitz  # pymupdf


def extract_pdf(pdf_path: Path) -> dict:
    """Extract text and TOC from a PDF file."""
    doc = fitz.open(str(pdf_path))

    # Extract TOC from bookmarks
    toc = []
    raw_toc = doc.get_toc(simple=True)  # [[level, title, page], ...]
    for level, title, page in raw_toc:
        toc.append({
            "title": title.strip(),
            "level": level,  # 1 = chapter, 2 = section, 3 = subsection
            "page": page,
        })

    # Extract text page by page
    pages = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages.append({
                "page": page_num + 1,  # 1-indexed
                "text": text,
            })

    metadata = {
        "title": doc.metadata.get("title", "") or pdf_path.stem,
        "author": doc.metadata.get("author", ""),
        "total_pages": len(doc),
    }

    doc.close()
    return {"pages": pages, "toc": toc, "metadata": metadata}


def extract_docx(docx_path: Path) -> dict:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(str(docx_path))

    # Build pseudo-pages (group paragraphs into ~500-word chunks)
    pages = []
    current_text = []
    current_word_count = 0
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            current_text.append("")
            continue

        current_text.append(text)
        current_word_count += len(text.split())

        if current_word_count >= 500:
            pages.append({
                "page": page_num,
                "text": "\n".join(current_text),
            })
            current_text = []
            current_word_count = 0
            page_num += 1

    if current_text:
        pages.append({
            "page": page_num,
            "text": "\n".join(current_text),
        })

    # Extract headings as TOC
    toc = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading", "").strip())
            except ValueError:
                level = 1
            toc.append({
                "title": para.text.strip(),
                "level": level,
                "page": 1,  # DOCX doesn't have page numbers easily
            })

    metadata = {
        "title": docx_path.stem,
        "author": "",
        "total_pages": page_num,
    }

    return {"pages": pages, "toc": toc, "metadata": metadata}


def extract(file_path: Path) -> dict:
    """Extract text from a PDF or DOCX file."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(file_path)
    elif suffix == ".docx":
        return extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python extract.py <input_file> [output_dir]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)

    from rich.console import Console
    console = Console()

    console.print(f"[bold blue]Extracting:[/] {input_path.name}")
    result = extract(input_path)

    output_file = output_dir / f"{input_path.stem}.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    console.print(
        f"  [green]✓[/] {len(result['pages'])} pages, "
        f"{len(result['toc'])} TOC entries → {output_file}"
    )


if __name__ == "__main__":
    main()
